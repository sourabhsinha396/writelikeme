from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from jose import JWTError, jwt
from datetime import datetime, timedelta
from typing import Optional
from sqlalchemy.orm import Session
from fastapi import Request, Form, File, UploadFile
from fastapi.responses import JSONResponse
from database.models import Sample, StyleProfile
from utils.style_analyzer import StyleAnalyzer
from database.database import get_db
from utils.auth import get_user_or_anonymous
from fastapi import Cookie
from typing import List
import os
import tempfile

router = APIRouter(tags=["samples"])


# @router.post("")
# async def add_sample_api(
#     request: Request,
#     profile_id: str = Form(...),
#     sample_text: str = Form(...),
#     sample_name: str = Form("Manual Entry"),
#     db: Session = Depends(get_db),
#     token: Optional[str] = Cookie(None, alias="access_token")
# ):
#     user, anonymous_user = get_user_or_anonymous(request, db, token)
    
#     # Check if profile exists
#     style_profile = db.query(StyleProfile).filter(StyleProfile.id == profile_id).first()
#     if not style_profile:
#         return JSONResponse(
#             status_code=404,
#             content={"error": "Profile not found"}
#         )
    
#     # Check ownership
#     profile_belongs_to_user = False
#     if user and style_profile.user_id == user.id:
#         profile_belongs_to_user = True
#     elif anonymous_user and style_profile.anonymous_user_id == anonymous_user.id:
#         profile_belongs_to_user = True
        
#     if not profile_belongs_to_user:
#         return JSONResponse(
#             status_code=403,
#             content={"error": "You don't have access to this profile"}
#         )
    
#     # Save the sample
#     sample = Sample(
#         style_profile_id=profile_id,
#         content=sample_text,
#         source_type="paste",
#         filename=sample_name
#     )
#     db.add(sample)
#     db.commit()
#     db.refresh(sample)
    
#     return JSONResponse(content={
#         "success": True,
#         "sample_id": sample.id,
#         "message": "Sample added successfully"
#     })

@router.delete("/delete-sample/{sample_id}")
async def delete_sample_api(
    request: Request,
    sample_id: str,
    db: Session = Depends(get_db),
    token: Optional[str] = Cookie(None, alias="access_token")
):
    user, anonymous_user = get_user_or_anonymous(request, db, token)
    
    # Get the sample
    sample = db.query(Sample).filter(Sample.id == sample_id).first()
    if not sample:
        return JSONResponse(
            status_code=404,
            content={"error": "Sample not found"}
        )
    
    # Get the profile to check ownership
    style_profile = db.query(StyleProfile).filter(StyleProfile.id == sample.style_profile_id).first()
    
    # Check ownership
    profile_belongs_to_user = False
    if user and style_profile.user_id == user.id:
        profile_belongs_to_user = True
    elif anonymous_user and style_profile.anonymous_user_id == anonymous_user.id:
        profile_belongs_to_user = True
        
    if not profile_belongs_to_user:
        return JSONResponse(
            status_code=403,
            content={"error": "You don't have access to this sample"}
        )
    
    # Delete the sample
    db.delete(sample)
    db.commit()
    
    return JSONResponse(content={"success": True})

@router.post("/retrain")
async def retrain_api(
    request: Request,
    profile_id: str = Form(...),
    db: Session = Depends(get_db),
    token: Optional[str] = Cookie(None, alias="access_token")
):
    user, anonymous_user = get_user_or_anonymous(request, db, token)
    
    # Get the profile
    style_profile = db.query(StyleProfile).filter(StyleProfile.id == profile_id).first()
    if not style_profile:
        return JSONResponse(
            status_code=404,
            content={"error": "Profile not found"}
        )
    
    # Check ownership
    profile_belongs_to_user = False
    if user and style_profile.user_id == user.id:
        profile_belongs_to_user = True
    elif anonymous_user and style_profile.anonymous_user_id == anonymous_user.id:
        profile_belongs_to_user = True
        
    if not profile_belongs_to_user:
        return JSONResponse(
            status_code=403,
            content={"error": "You don't have access to this profile"}
        )
    
    samples = db.query(Sample).filter(Sample.style_profile_id == profile_id).all()
    if not samples:
        return JSONResponse(
            status_code=400,
            content={"error": "No samples found to retrain on"}
        )
    
    # Initialize analyzer
    analyzer = StyleAnalyzer()
    
    # Add samples to analyzer
    for sample in samples:
        analyzer.add_sample(sample.content)
    
    # Analyze the style
    try:
        style_profile_data = analyzer.analyze()
        
        # Update the style profile with new analysis data
        style_profile.profile_data = style_profile_data
        db.commit()
        
        return JSONResponse(content={
            "success": True,
            "message": "Style profile retrained successfully"
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error retraining style profile: {str(e)}"}
        )
    

@router.post("/upload")
async def upload_samples_api(
    request: Request,
    name: str = Form(...),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    token: Optional[str] = Cookie(None, alias="access_token")
):
    user, anonymous_user = get_user_or_anonymous(request, db, token)
    # get preexisting style profile
    style_profile = db.query(StyleProfile).filter(StyleProfile.name == name).first()
    if not style_profile:
        style_profile = StyleProfile(
            name=name
        )
        db.add(style_profile)
        db.commit()
        db.refresh(style_profile)
    
    if user:
        style_profile.user_id = user.id
    else:
        style_profile.anonymous_user_id = anonymous_user.id
    
    analyzer = StyleAnalyzer()
    successful_samples = 0
    error_files = []
    
    for file in files:
        try:
            print(f"Processing file: {file.filename}")
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp:
                # Read content in chunks to handle large files
                content = await file.read()
                temp.write(content)
                temp_path = temp.name
            
            text = ""
            file_extension = os.path.splitext(file.filename)[1].lower()
            
            # Extract text based on file type
            if file_extension == '.txt':
                with open(temp_path, 'r', errors='ignore') as f:
                    text = f.read()
                print(f"TXT file processed, extracted {len(text)} characters")
            
            elif file_extension == '.docx':
                try:
                    import docx
                    doc = docx.Document(temp_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    print(f"DOCX file processed, extracted {len(text)} characters from {len(doc.paragraphs)} paragraphs")
                except Exception as doc_err:
                    print(f"Error processing DOCX: {str(doc_err)}")
                    error_files.append(f"{file.filename} (DOCX error: {str(doc_err)})")
            
            elif file_extension == '.pdf':
                try:
                    import PyPDF2
                    with open(temp_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text_list = []
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                text_list.append(page_text)
                        text = "\n".join(text_list)
                        print(f"PDF text extraction attempt 1: extracted {len(text)} characters from {len(pdf_reader.pages)} pages")
                    
                    # If no text was extracted, try OCR
                    if not text.strip():
                        print("No text extracted from PDF using PyPDF2, trying OCR...")
                        try:
                            import pytesseract
                            from pdf2image import convert_from_path
                            
                            # Convert PDF to images
                            images = convert_from_path(temp_path)
                            ocr_text_list = []
                            
                            for i, image in enumerate(images):
                                # Use pytesseract to extract text from image
                                page_text = pytesseract.image_to_string(image)
                                ocr_text_list.append(page_text)
                            
                            text = "\n".join(ocr_text_list)
                            print(f"PDF OCR successful: extracted {len(text)} characters from {len(images)} pages")
                        except Exception as ocr_err:
                            print(f"Error in OCR process: {str(ocr_err)}")
                            error_files.append(f"{file.filename} (OCR error: {str(ocr_err)})")
                
                except Exception as pdf_err:
                    print(f"Error processing PDF: {str(pdf_err)}")
                    error_files.append(f"{file.filename} (PDF error: {str(pdf_err)})")
            
            else:
                print(f"Unknown file type: {file_extension}")
                with open(temp_path, 'r', errors='ignore') as f:
                    text = f.read()
                print(f"Attempted to read as text, extracted {len(text)} characters")
            
            # Only add sample if text was successfully extracted
            if text and text.strip():
                analyzer.add_sample(text)
                successful_samples += 1
                
                # Save the sample
                sample = Sample(
                    style_profile_id=style_profile.id,
                    content=text,
                    source_type="upload",
                    filename=file.filename
                )
                db.add(sample)
                print(f"Sample added for {file.filename}")
            else:
                print(f"No valid text extracted from {file.filename}")
                error_files.append(f"{file.filename} (No text extracted)")
            
            os.unlink(temp_path)
            
        except Exception as e:
            print(f"General error processing {file.filename}: {str(e)}")
            error_files.append(f"{file.filename} (Error: {str(e)})")
            continue
    
    db.commit()
    
    if successful_samples == 0:
        return JSONResponse(
            status_code=400,
            content={
                "error": "No valid text could be extracted from the uploaded files",
                "failed_files": error_files
            }
        )
    
    try:
        print(f"Analyzing {successful_samples} samples")
        style_profile_data = analyzer.analyze()
        
        # Update the style profile with analysis data
        style_profile.profile_data = style_profile_data
        db.commit()
        
        return JSONResponse(content={
            "success": True,
            "profile_id": style_profile.id,
            "message": f"Files uploaded and analyzed successfully ({successful_samples} samples processed)",
            "warnings": error_files if error_files else None
        })
    except Exception as e:
        print(f"Error in analyzing style: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Error analyzing style: {str(e)}"}
        )
    

@router.post("/analyze-text")
async def analyze_text_api(request: Request, sample_name: str = Form(...), sample_text: str = Form(...), db: Session = Depends(get_db),
    token: Optional[str] = Cookie(None, alias="access_token")
):
    user, anonymous_user = get_user_or_anonymous(request, db, token)
    
    style_profile = StyleProfile(
        name=sample_name
    )
    
    if user:
        style_profile.user_id = user.id
    else:
        style_profile.anonymous_user_id = anonymous_user.id
    
    db.add(style_profile)
    db.commit()
    db.refresh(style_profile)
    
    # Initialize analyzer
    analyzer = StyleAnalyzer()
    analyzer.add_sample(sample_text)
    
    # Save the sample
    sample = Sample(
        style_profile_id=style_profile.id,
        content=sample_text,
        source_type="paste",
        filename=sample_name
    )
    db.add(sample)
    db.commit()
    
    # Analyze the style
    try:
        style_profile_data = analyzer.analyze()
        style_profile.profile_data = style_profile_data
        db.commit()
        
        return JSONResponse(content={
            "success": True,
            "profile_id": style_profile.id,
            "message": "Text analyzed successfully"
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error analyzing style: {str(e)}"}
        )